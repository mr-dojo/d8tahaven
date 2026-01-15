"""File text extraction utilities for PDF, DOCX, and TXT files."""

import io
import structlog
from pypdf import PdfReader
from docx import Document
import chardet

logger = structlog.get_logger()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


class ExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text content from a PDF file.

    Args:
        file_content: Raw bytes of the PDF file

    Returns:
        Extracted text as a string

    Raises:
        ExtractionError: If PDF parsing fails
    """
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(text)

        result = "\n".join(text_parts)
        logger.info("pdf_extracted", pages=len(reader.pages), text_length=len(result))
        return result
    except Exception as e:
        logger.error("pdf_extraction_failed", error=str(e))
        raise ExtractionError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from a DOCX file.

    Args:
        file_content: Raw bytes of the DOCX file

    Returns:
        Extracted text as a string

    Raises:
        ExtractionError: If DOCX parsing fails
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)

        result = "\n".join(text_parts)
        logger.info("docx_extracted", paragraphs=len(doc.paragraphs), text_length=len(result))
        return result
    except Exception as e:
        logger.error("docx_extraction_failed", error=str(e))
        raise ExtractionError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text content from a TXT file with automatic encoding detection.

    Uses chardet to detect encoding, falls back to common encodings if detection fails.

    Args:
        file_content: Raw bytes of the TXT file

    Returns:
        Extracted text as a string

    Raises:
        ExtractionError: If decoding fails with all attempted encodings
    """
    # Try to detect encoding
    detection = chardet.detect(file_content)
    detected_encoding = detection.get('encoding')
    confidence = detection.get('confidence', 0)

    logger.debug(
        "encoding_detected",
        encoding=detected_encoding,
        confidence=confidence
    )

    # Build list of encodings to try
    encodings_to_try = []

    # If chardet is confident, try that first
    if detected_encoding and confidence > 0.7:
        encodings_to_try.append(detected_encoding)

    # Add common encodings as fallbacks
    for enc in ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']:
        if enc.lower() not in [e.lower() for e in encodings_to_try]:
            encodings_to_try.append(enc)

    # Try each encoding
    for encoding in encodings_to_try:
        try:
            result = file_content.decode(encoding)
            logger.info("txt_extracted", encoding=encoding, text_length=len(result))
            return result
        except (UnicodeDecodeError, LookupError) as e:
            logger.debug("encoding_failed", encoding=encoding, error=str(e))
            continue

    # If all fail, raise error
    raise ExtractionError(
        f"Failed to decode text file. Tried encodings: {', '.join(encodings_to_try)}. "
        f"Detected encoding was: {detected_encoding} (confidence: {confidence:.2f})"
    )


def extract_text(filename: str, file_content: bytes) -> str:
    """
    Extract text from file based on extension.

    Args:
        filename: Original filename (used to determine file type)
        file_content: Raw bytes of the file

    Returns:
        Extracted text as a string

    Raises:
        ValueError: If file type is not supported
        ExtractionError: If extraction fails
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == ".pdf":
        return extract_text_from_pdf(file_content)
    elif ext == ".docx":
        return extract_text_from_docx(file_content)
    elif ext == ".txt":
        return extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
